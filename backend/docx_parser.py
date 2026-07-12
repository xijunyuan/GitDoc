"""
DocxParser
==========
Extracts structured text content from .docx files by parsing the
underlying OOXML (word/document.xml).

Uses python-docx to iterate over paragraphs, tables, and lists in
document order. Handles nested tables within table cells.
"""

import json
from pathlib import Path
from typing import List, Optional, Dict, Any, Generator, Union

try:
    from docx import Document
    from docx.document import Document as _Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph
except ModuleNotFoundError:
    import sys
    raise ImportError(
        f"python-docx not installed for {sys.executable}.\n"
        f"Run: pip install python-docx"
    ) from None


class Block:
    """
    A single content block extracted from a .docx document.
    Blocks represent paragraphs, tables, or list items in document order.
    """

    def __init__(self, block_type: str, text: str, index: int,
                 style: Optional[str] = None, extra: Optional[Dict[str, Any]] = None):
        self.type = block_type       # "paragraph" | "table" | "list"
        self.text = text
        self.index = index
        self.style = style
        self.extra = extra or {}

    def to_dict(self) -> dict:
        return {
            "type": self.type,
            "text": self.text,
            "index": self.index,
            "style": self.style
        }


class DocumentContent:
    """
    Structured representation of a .docx file's extractable content.
    Holds an ordered list of Block objects and provides the full text.
    """

    def __init__(self, blocks: List[Block]):
        self.blocks = blocks

    @property
    def full_text(self) -> str:
        """Concatenated plain text with double-newline separators between blocks."""
        return "\n\n".join(b.text for b in self.blocks)

    def to_cache_dict(self) -> dict:
        """Serialize to a cacheable dictionary."""
        return {
            "blocks": [b.to_dict() for b in self.blocks],
            "full_text": self.full_text
        }


class DocxParser:
    """
    Extracts structured content from .docx files.

    Handles paragraphs, tables (including nested tables), and preserves
    document order as the content appears in the OOXML body.
    """

    @staticmethod
    def iter_block_items(parent) -> Generator[Union[Paragraph, Table], None, None]:
        """
        Yield Paragraph and Table objects in document order.

        This is the canonical pattern from python-docx documentation.
        Walks the children of <w:body> (or a table cell's <w:tc>) and
        yields python-docx Paragraph or Table objects based on the XML tag.
        """
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError(f"Unsupported parent type: {type(parent)}")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    @staticmethod
    def _extract_cell_text(cell) -> str:
        """
        Extract text from a table cell, handling nested tables.
        For cells containing nested tables, flattens the table content.
        """
        # Check for nested tables
        if cell.tables:
            rows = []
            for nested_table in cell.tables:
                for row in nested_table.rows:
                    cell_texts = [c.text.strip() for c in row.cells]
                    rows.append(" | ".join(cell_texts))
            return "\n".join(rows)
        return cell.text.strip()

    @staticmethod
    def extract_blocks(docx_path: str) -> List[Block]:
        """
        Extract all content blocks from a .docx file, preserving document order.

        Args:
            docx_path: Path to the .docx file.

        Returns:
            Ordered list of Block objects (paragraphs, tables, etc.).
        """
        doc = Document(docx_path)
        blocks = []
        index = 0

        for block_item in DocxParser.iter_block_items(doc):
            if isinstance(block_item, Paragraph):
                text = block_item.text.strip()
                style = None
                try:
                    style = block_item.style.name if block_item.style else None
                except Exception:
                    style = None

                if text:  # Skip truly empty paragraphs
                    # Detect list items
                    pPr = block_item._element.find(
                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr'
                    )
                    if pPr is not None:
                        numPr = pPr.find(
                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr'
                        )
                        if numPr is not None:
                            blocks.append(Block("list", text, index, style=style))
                            index += 1
                            continue

                    blocks.append(Block("paragraph", text, index, style=style))
                    index += 1

            elif isinstance(block_item, Table):
                rows_text = []
                for row in block_item.rows:
                    cells = [DocxParser._extract_cell_text(cell) for cell in row.cells]
                    rows_text.append(" | ".join(cells))
                table_text = "\n".join(rows_text)

                if table_text.strip():
                    blocks.append(Block(
                        "table", table_text, index,
                        extra={
                            "rows": len(block_item.rows),
                            "cols": len(block_item.columns)
                        }
                    ))
                    index += 1

        return blocks

    @staticmethod
    def extract_text(docx_path: str) -> str:
        """
        Quick plain-text extraction without full block structuring.
        Useful when only the full text is needed (e.g., simple preview).
        """
        blocks = DocxParser.extract_blocks(docx_path)
        return "\n\n".join(b.text for b in blocks)

    # ---- Caching ----

    @staticmethod
    def extract_and_cache(docx_path: str, commit_hash: str, cache_dir: Path) -> str:
        """
        Extract blocks from a .docx, cache them as JSON, and return the full text.
        Creates the cache directory if it doesn't exist.
        """
        content = DocumentContent(DocxParser.extract_blocks(docx_path))
        cache_dir.mkdir(parents=True, exist_ok=True)
        cache_path = cache_dir / f"{commit_hash}.json"
        with open(cache_path, "w", encoding="utf-8") as f:
            json.dump(content.to_cache_dict(), f, ensure_ascii=False, indent=2)
        return content.full_text

    @staticmethod
    def load_cached_text(commit_hash: str, cache_dir: Path) -> Optional[str]:
        """
        Load previously cached text for a commit hash.
        Returns None if the cache file doesn't exist.
        """
        cache_path = cache_dir / f"{commit_hash}.json"
        if not cache_path.exists():
            return None
        with open(cache_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data.get("full_text")

    @staticmethod
    def load_cached_blocks(commit_hash: str, cache_dir: Path) -> Optional[List[Block]]:
        """
        Load previously cached blocks for a commit hash.
        Returns None if the cache file doesn't exist.
        """
        cache_path = cache_dir / f"{commit_hash}.json"
        if not cache_path.exists():
            return None
        with open(cache_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        blocks_data = data.get("blocks", [])
        blocks = []
        for i, b in enumerate(blocks_data):
            blocks.append(Block(
                block_type=b.get("type", "paragraph"),
                text=b.get("text", ""),
                index=b.get("index", i),
                style=b.get("style")
            ))
        return blocks

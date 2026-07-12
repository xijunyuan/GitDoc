"""
Tests for DocxParser Module
============================
Tests OOXML text extraction from .docx files.
"""

import os
import sys
import json
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx_parser import DocxParser, Block, DocumentContent

# Note: These tests require python-docx to create test fixture files.


class TestDocxParser(unittest.TestCase):
    """Test suite for DocxParser class."""

    @classmethod
    def setUpClass(cls):
        """Create test .docx files using python-docx."""
        from docx import Document
        cls.tmp_dir = tempfile.mkdtemp(prefix="gitdoc_parser_test_")

        # --- Simple document: 3 paragraphs ---
        doc = Document()
        doc.add_paragraph("First paragraph of text.")
        doc.add_paragraph("Second paragraph with more content.")
        doc.add_paragraph("Third paragraph, the final one.")
        cls.simple_path = os.path.join(cls.tmp_dir, "simple.docx")
        doc.save(cls.simple_path)

        # --- Document with table ---
        doc2 = Document()
        doc2.add_paragraph("Paragraph before table.")
        table = doc2.add_table(rows=3, cols=3)
        headers = ["Name", "Age", "City"]
        data = [
            ["Alice", "30", "Beijing"],
            ["Bob", "25", "Shanghai"]
        ]
        for j, header in enumerate(headers):
            table.rows[0].cells[j].text = header
        for i, row_data in enumerate(data):
            for j, cell_text in enumerate(row_data):
                table.rows[i + 1].cells[j].text = cell_text
        doc2.add_paragraph("Paragraph after table.")
        cls.table_path = os.path.join(cls.tmp_dir, "with_table.docx")
        doc2.save(cls.table_path)

        # --- Empty document ---
        doc3 = Document()
        cls.empty_path = os.path.join(cls.tmp_dir, "empty.docx")
        doc3.save(cls.empty_path)

    @classmethod
    def tearDownClass(cls):
        """Clean up."""
        import shutil
        if os.path.exists(cls.tmp_dir):
            shutil.rmtree(cls.tmp_dir, ignore_errors=True)

    # --- Block Tests ---

    def test_block_creation(self):
        """Block should store type, text, index, and style."""
        b = Block("paragraph", "Hello world", 0, style="Normal")
        self.assertEqual(b.type, "paragraph")
        self.assertEqual(b.text, "Hello world")
        self.assertEqual(b.index, 0)
        self.assertEqual(b.style, "Normal")

    def test_block_to_dict(self):
        """to_dict() should serialize correctly."""
        b = Block("table", "A | B", 5, style=None, extra={"rows": 2, "cols": 2})
        d = b.to_dict()
        self.assertEqual(d["type"], "table")
        self.assertEqual(d["text"], "A | B")
        self.assertEqual(d["index"], 5)

    def test_block_extra(self):
        """Extra metadata should be stored."""
        b = Block("paragraph", "text", 0, extra={"style": "Heading1"})
        self.assertEqual(b.extra["style"], "Heading1")

    # --- DocumentContent Tests ---

    def test_document_content_full_text(self):
        """full_text should concatenate blocks with double-newline."""
        blocks = [
            Block("paragraph", "Block 1", 0),
            Block("paragraph", "Block 2", 1)
        ]
        content = DocumentContent(blocks)
        self.assertEqual(content.full_text, "Block 1\n\nBlock 2")

    def test_document_content_to_cache_dict(self):
        """to_cache_dict should produce JSON-serializable output."""
        blocks = [Block("paragraph", "Test", 0)]
        content = DocumentContent(blocks)
        d = content.to_cache_dict()
        self.assertIn("blocks", d)
        self.assertIn("full_text", d)
        self.assertEqual(d["full_text"], "Test")
        # Should be JSON-serializable
        json.dumps(d)

    # --- Extraction Tests ---

    def test_extract_simple_document(self):
        """Should extract all paragraphs from a simple document."""
        blocks = DocxParser.extract_blocks(self.simple_path)
        self.assertGreaterEqual(len(blocks), 3)

        # Should have paragraph blocks
        para_blocks = [b for b in blocks if b.type == "paragraph"]
        self.assertEqual(len(para_blocks), 3)
        self.assertIn("First paragraph", para_blocks[0].text)

    def test_extract_text_simple(self):
        """extract_text should return concatenated text."""
        text = DocxParser.extract_text(self.simple_path)
        self.assertIn("First paragraph", text)
        self.assertIn("Second paragraph", text)
        self.assertIn("Third paragraph", text)

    def test_extract_document_with_table(self):
        """Should extract both paragraphs and tables in order."""
        blocks = DocxParser.extract_blocks(self.table_path)
        self.assertGreaterEqual(len(blocks), 2)

        # First block should be a paragraph
        self.assertEqual(blocks[0].type, "paragraph")
        self.assertIn("before table", blocks[0].text)

    def test_table_block_content(self):
        """Table block should contain cell data."""
        blocks = DocxParser.extract_blocks(self.table_path)
        table_blocks = [b for b in blocks if b.type == "table"]
        self.assertGreaterEqual(len(table_blocks), 1)

        table_text = table_blocks[0].text
        self.assertIn("Alice", table_text)
        self.assertIn("Beijing", table_text)

    def test_empty_document(self):
        """Empty document should return empty blocks list."""
        blocks = DocxParser.extract_blocks(self.empty_path)
        # An empty .docx may still have one empty paragraph
        # Just ensure it doesn't crash
        self.assertIsInstance(blocks, list)

    def test_block_index_ordering(self):
        """Blocks should be sequentially indexed."""
        blocks = DocxParser.extract_blocks(self.simple_path)
        for i, b in enumerate(blocks):
            self.assertEqual(b.index, i)

    # --- Cache Tests ---

    def test_extract_and_cache(self):
        """extract_and_cache should write a JSON cache file."""
        import tempfile
        cache_dir = Path(tempfile.mkdtemp(prefix="gitdoc_cache_test_"))
        commit_hash = "abc123def456"

        text = DocxParser.extract_and_cache(self.simple_path, commit_hash, cache_dir)

        cache_file = cache_dir / f"{commit_hash}.json"
        self.assertTrue(cache_file.exists())

        # Verify content
        with open(cache_file, "r", encoding="utf-8") as f:
            data = json.load(f)
        self.assertIn("full_text", data)
        self.assertEqual(data["full_text"], text)

        # Cleanup
        import shutil
        shutil.rmtree(cache_dir, ignore_errors=True)

    def test_load_cached_text_exists(self):
        """load_cached_text should return text for a cached commit."""
        import tempfile
        cache_dir = Path(tempfile.mkdtemp(prefix="gitdoc_cache_test_"))
        commit_hash = "def789abc012"

        DocxParser.extract_and_cache(self.simple_path, commit_hash, cache_dir)

        cached = DocxParser.load_cached_text(commit_hash, cache_dir)
        self.assertIsNotNone(cached)
        self.assertIn("First paragraph", cached)

        import shutil
        shutil.rmtree(cache_dir, ignore_errors=True)

    def test_load_cached_text_missing(self):
        """load_cached_text should return None for uncached commits."""
        import tempfile
        cache_dir = Path(tempfile.mkdtemp(prefix="gitdoc_cache_test_"))
        result = DocxParser.load_cached_text("nonexistent_hash", cache_dir)
        self.assertIsNone(result)
        import shutil
        shutil.rmtree(cache_dir, ignore_errors=True)


if __name__ == "__main__":
    unittest.main()

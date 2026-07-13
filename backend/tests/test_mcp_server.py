"""
Tests for GitDoc MCP Server
===========================
Tests the tool implementation functions (_*_impl) directly.
These are the sync workers invoked by the async tool wrappers.

Error handling (ValueError → error dict) is tested via the _safe_call
wrapper in mcp_server.py, which is verified through the MCP client
integration test.
"""

import sys
import os
import tempfile
from pathlib import Path

import pytest

# Ensure backend/ is on sys.path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document as DocxDocument

# Tool implementations (sync — test these directly)
from mcp_server import (
    _read_docx_impl,
    _get_history_impl,
    _diff_versions_impl,
    _preview_version_impl,
    _diff_files_impl,
    _validate_docx_path,
    _require_repo,
    _resolve_hash,
    _load_content_at_commit,
    _parse_docx_bytes,
)
from git_operations import GitRepo


# ============================================================
# Fixtures
# ============================================================

def _make_docx(dir_path: Path, filename: str, paragraphs: list[str]) -> str:
    """Create a .docx file with the given paragraphs and return its path."""
    filepath = str(dir_path / filename)
    doc = DocxDocument()
    for text in paragraphs:
        if text:
            doc.add_paragraph(text)
    doc.save(filepath)
    return filepath


@pytest.fixture
def work_dir():
    """Isolated temp directory — cleaned up after test."""
    with tempfile.TemporaryDirectory() as td:
        yield Path(td)


@pytest.fixture
def sample_docx(work_dir):
    """A .docx file with known content in an isolated temp directory."""
    yield _make_docx(work_dir, "sample.docx", [
        "Hello world. This is the first paragraph.",
        "Second paragraph with more content.",
        "",  # empty — will be skipped by parser
    ])


@pytest.fixture
def sample_docx_v2(work_dir):
    """A second .docx with modified content (for diff tests)."""
    yield _make_docx(work_dir, "sample_v2.docx", [
        "Hello world. This is the first paragraph.",
        "Second paragraph has been CHANGED here.",
        "A brand new third paragraph.",
    ])


@pytest.fixture
def tracked_docx(work_dir):
    """A .docx file with an initialized GitDoc repo and two commits.

    Lives in its own temp directory, fully isolated from other fixtures.
    """
    path = _make_docx(work_dir, "tracked.docx", [
        "Hello world. This is the first paragraph.",
        "Second paragraph with more content.",
    ])

    repo = GitRepo(path)
    repo.init()
    repo.commit("[test] v1 — initial version")

    # Make a second version by modifying the doc
    doc = DocxDocument()
    doc.add_paragraph("Hello world. This is the first paragraph — MODIFIED.")
    doc.add_paragraph("Second paragraph with more content.")
    doc.add_paragraph("A completely new paragraph added in v2.")
    doc.save(path)
    repo.commit("[test] v2 — modified and added paragraph")

    yield path, repo


# ============================================================
# read_docx
# ============================================================

class TestReadDocx:
    """Tests for the read_docx implementation."""

    def test_parses_valid_docx(self, sample_docx):
        result = _read_docx_impl(sample_docx)
        assert result["block_count"] >= 2
        assert "Hello world" in result["full_text"]
        assert result["path"] == sample_docx
        assert isinstance(result["blocks"], list)
        assert "type" in result["blocks"][0]

    def test_returns_structured_blocks(self, sample_docx):
        result = _read_docx_impl(sample_docx)
        types = {b["type"] for b in result["blocks"]}
        assert "paragraph" in types

    def test_file_not_found(self):
        with pytest.raises(ValueError, match="File not found"):
            _read_docx_impl("f:/nonexistent_file_12345.docx")

    def test_empty_path(self):
        with pytest.raises(ValueError, match="Path cannot be empty"):
            _read_docx_impl("")

    def test_non_docx_file(self, tmp_path):
        txt = tmp_path / "readme.txt"
        txt.write_text("hello")
        with pytest.raises(ValueError, match="Not a Word document"):
            _read_docx_impl(str(txt))


# ============================================================
# get_history
# ============================================================

class TestGetHistory:
    """Tests for the get_history implementation."""

    def test_returns_commits_for_tracked_docx(self, tracked_docx):
        path, _repo = tracked_docx
        result = _get_history_impl(path, 50)
        assert result["total_count"] >= 2
        assert len(result["commits"]) >= 2
        assert "hash" in result["commits"][0]
        assert "short_hash" in result["commits"][0]
        assert "message" in result["commits"][0]
        assert "timestamp" in result["commits"][0]

    def test_max_count_respected(self, tracked_docx):
        path, _repo = tracked_docx
        result = _get_history_impl(path, 1)
        assert result["total_count"] == 1

    def test_no_repo_raises(self, work_dir):
        """A file in a clean directory with no .gitdoc/ should fail."""
        path = _make_docx(work_dir, "fresh.docx", ["test"])
        with pytest.raises(ValueError, match="No GitDoc repository"):
            _get_history_impl(path, 50)


# ============================================================
# diff_versions
# ============================================================

class TestDiffVersions:
    """Tests for the diff_versions implementation."""

    def test_detects_changes(self, tracked_docx):
        path, repo = tracked_docx
        commits = repo.get_history(max_count=10)
        v1_hash = commits[-1].hash  # oldest
        v2_hash = commits[0].hash   # newest

        result = _diff_versions_impl(path, v1_hash, v2_hash)
        assert "blocks" in result
        assert "stats" in result
        assert result["from_hash"] == v1_hash
        assert result["to_hash"] == v2_hash
        stats = result["stats"]
        assert stats["insertions"] > 0 or stats["deletions"] > 0

    def test_same_version_no_diff(self, tracked_docx):
        path, repo = tracked_docx
        commits = repo.get_history(max_count=10)
        h = commits[0].hash

        result = _diff_versions_impl(path, h, h)
        stats = result["stats"]
        assert stats["insertions"] == 0
        assert stats["deletions"] == 0

    def test_short_hash_resolution(self, tracked_docx):
        path, repo = tracked_docx
        commits = repo.get_history(max_count=10)
        short1 = commits[-1].short_hash
        short2 = commits[0].short_hash

        result = _diff_versions_impl(path, short1, short2)
        assert len(result["from_hash"]) == 40

    def test_invalid_hash(self, tracked_docx):
        path, _repo = tracked_docx
        with pytest.raises(ValueError, match="No commit found"):
            _diff_versions_impl(path, "00000000", "11111111")

    def test_no_repo_raises(self, work_dir):
        path = _make_docx(work_dir, "fresh.docx", ["test"])
        with pytest.raises(ValueError, match="No GitDoc repository"):
            _diff_versions_impl(path, "abc", "def")


# ============================================================
# preview_version
# ============================================================

class TestPreviewVersion:
    """Tests for the preview_version implementation."""

    def test_returns_text_for_version(self, tracked_docx):
        path, repo = tracked_docx
        commits = repo.get_history(max_count=10)
        h = commits[0].hash

        result = _preview_version_impl(path, h)
        assert result["hash"] == h
        assert isinstance(result["text"], str)
        assert result["block_count"] >= 1

    def test_short_hash_resolves(self, tracked_docx):
        path, repo = tracked_docx
        sh = repo.get_history(max_count=1)[0].short_hash

        result = _preview_version_impl(path, sh)
        assert len(result["hash"]) == 40

    def test_invalid_hash_raises(self, tracked_docx):
        path, _repo = tracked_docx
        with pytest.raises(ValueError, match="No commit found"):
            _preview_version_impl(path, "feeddead")


# ============================================================
# diff_files
# ============================================================

class TestDiffFiles:
    """Tests for the diff_files implementation."""

    def test_identical_files_no_diff(self, sample_docx):
        result = _diff_files_impl(sample_docx, sample_docx)
        assert result["file1"] == sample_docx
        assert result["file2"] == sample_docx
        assert result["stats"]["insertions"] == 0
        assert result["stats"]["deletions"] == 0

    def test_different_files(self, sample_docx, sample_docx_v2):
        result = _diff_files_impl(sample_docx, sample_docx_v2)
        assert "blocks" in result
        stats = result["stats"]
        assert stats["insertions"] > 0 or stats["deletions"] > 0

    def test_file1_not_found(self):
        with pytest.raises(ValueError, match="File not found"):
            _diff_files_impl("f:/ghost.docx", "f:/also_ghost.docx")

    def test_file2_not_found(self, sample_docx):
        with pytest.raises(ValueError, match="File not found"):
            _diff_files_impl(sample_docx, "f:/ghost.docx")

    def test_non_docx_rejected(self, sample_docx, tmp_path):
        txt = tmp_path / "readme.txt"
        txt.write_text("hello")
        with pytest.raises(ValueError, match="Word document"):
            _diff_files_impl(sample_docx, str(txt))


# ============================================================
# Internal helpers
# ============================================================

class TestInternalHelpers:
    """Tests for the non-tool helper functions."""

    def test_validate_docx_path_accepts_doc_as_well(self, tmp_path):
        """_validate_docx_path should accept both .docx and .doc extensions."""
        doc_file = tmp_path / "legacy.doc"
        doc_file.write_text("fake doc")
        _validate_docx_path(str(doc_file))  # should NOT raise

    def test_require_repo_when_not_initialized(self, work_dir):
        path = _make_docx(work_dir, "fresh.docx", ["no repo"])
        with pytest.raises(ValueError, match="No GitDoc repository"):
            _require_repo(path)

    def test_require_repo_when_initialized(self, tracked_docx):
        path, _repo = tracked_docx
        repo = _require_repo(path)
        assert repo.is_initialized()

    def test_parse_docx_bytes(self):
        """_parse_docx_bytes should work with in-memory .docx bytes."""
        import io
        doc = DocxDocument()
        doc.add_paragraph("Test from bytes")
        buf = io.BytesIO()
        doc.save(buf)
        raw = buf.getvalue()

        content = _parse_docx_bytes(raw)
        assert "Test from bytes" in content.full_text
        assert content.blocks[0].text == "Test from bytes"

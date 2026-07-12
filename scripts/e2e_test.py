"""
GitDoc End-to-End Test Script
==============================
Runs a full E2E test against a running backend.
Usage: python e2e_test.py
"""
import sys
import os
import json
import ssl
import urllib.request
import urllib.error
import urllib.parse
import tempfile

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'backend'))
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx_parser import DocxParser

# Try HTTPS first, fall back to HTTP
BASE = 'https://localhost:18521'
_ssl_ctx = ssl.create_default_context()
_ssl_ctx.check_hostname = False
_ssl_ctx.verify_mode = ssl.CERT_NONE

errors = []


def api(method, path, data=None):
    url = BASE + path
    body = json.dumps(data).encode() if data else None
    req = urllib.request.Request(url, data=body, method=method)
    req.add_header('Content-Type', 'application/json')
    try:
        with urllib.request.urlopen(req, context=_ssl_ctx) as r:
            return json.loads(r.read())
    except urllib.error.HTTPError as e:
        return {'error': e.code, 'body': e.read().decode()[:300]}
    except (urllib.error.URLError, ConnectionRefusedError, OSError):
        # Fallback: try HTTP
        url = 'http://127.0.0.1:18521' + path
        req = urllib.request.Request(url, data=body, method=method)
        req.add_header('Content-Type', 'application/json')
        try:
            with urllib.request.urlopen(req) as r:
                return json.loads(r.read())
        except urllib.error.HTTPError as e:
            return {'error': e.code, 'body': e.read().decode()[:300]}


def check(desc, condition):
    if condition:
        print(f"  [OK] {desc}")
    else:
        print(f"  [FAIL] {desc}")
        errors.append(desc)


# 1. Status
r = api('GET', '/api/status')
check(f"Status: {r.get('git_version', '?')}", r.get('status') == 'ok')

# 2. Create test docx
tmp = tempfile.mkdtemp(prefix='gitdoc_demo_')
docx = os.path.join(tmp, 'demo.docx')
doc = Document()
doc.add_paragraph('第一章：引言')
doc.add_paragraph('本文档用于演示 GitDoc 版本管理功能。')
doc.add_paragraph('包含三个段落。')
doc.save(docx)
check(f"Created test docx", os.path.exists(docx))

# 3. Init
r = api('POST', '/api/init', {'docx_path': docx, 'author': 'Alice'})
check("Init repo", r.get('success'))

# 4. Commit v1
r = api('POST', '/api/commit', {'docx_path': docx, 'message': '[manual] 初始版本', 'author': 'Alice'})
if r.get('success'):
    h1 = r['hash']
    check(f"Commit v1: {h1[:8]}", True)
else:
    check(f"Commit v1: {r}", False)

# 5. No-change commit rejected
r = api('POST', '/api/commit', {'docx_path': docx, 'message': 'No change'})
check("No-change rejected", not r.get('success'))

# 6. Commit v2
doc = Document(docx)
doc.add_paragraph('第四节：新增加的内容。')
doc.save(docx)
r = api('POST', '/api/commit', {'docx_path': docx, 'message': '[manual] 添加章节', 'author': 'Bob'})
if r.get('success'):
    h2 = r['hash']
    check(f"Commit v2: {h2[:8]}", True)
else:
    check(f"Commit v2: {r}", False)

# 7. History
r = api('GET', f'/api/history?docx_path={urllib.parse.quote(docx)}')
commits = r.get('commits', [])
check(f"History: {len(commits)} commits", len(commits) == 2)

# 8. Diff
r = api('GET', f'/api/diff?from_hash={h1}&to_hash={h2}&docx_path={urllib.parse.quote(docx)}')
stats = r.get('stats', {})
check(f"Diff: +{stats.get('insertions', 0)} insertions", stats.get('insertions', 0) > 0)

# 9. Preview
r = api('GET', f'/api/preview?commit_hash={h1}&docx_path={urllib.parse.quote(docx)}')
check(f"Preview: {r.get('block_count', 0)} blocks", bool(r.get('text')))

# 10. Rollback
r = api('POST', '/api/rollback', {'commit_hash': h1, 'docx_path': docx})
if r.get('success'):
    text = DocxParser.extract_text(docx)
    check("Rollback content verified", '新增加' not in text)
else:
    check(f"Rollback: {r}", False)

print()
if not errors:
    print("=" * 50)
    print("ALL 10 E2E TESTS PASSED")
    print("=" * 50)
    sys.exit(0)
else:
    print("=" * 50)
    print(f"{len(errors)} TESTS FAILED")
    print("=" * 50)
    sys.exit(1)
